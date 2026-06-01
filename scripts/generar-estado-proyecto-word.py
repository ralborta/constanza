#!/usr/bin/env python3
"""Genera documento Word con el estado actual del proyecto Constanza.

Incluye: qué hay implementado, cómo está funcionando, y qué falta para llegar
al 100% del alcance comprometido hasta este punto.
"""

from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


# ---------- helpers ----------

def add_title(doc, text, size=24, color=(15, 23, 42), align="center"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


def add_subtitle(doc, text, italic=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 80, 80)
    return p


def add_h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(15, 23, 42)
    return h


def add_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = RGBColor(30, 41, 59)
    return h


def add_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = RGBColor(51, 65, 85)
    return h


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else f"List Bullet {min(level + 1, 3)}"
    try:
        doc.add_paragraph(text, style=style)
    except KeyError:
        doc.add_paragraph(f"- {text}")


def add_status_table(doc, rows):
    """rows: list of (componente, estado, detalle)."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Componente"
    hdr[1].text = "Estado"
    hdr[2].text = "Detalle / Observación"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (comp, st, det) in enumerate(rows, start=1):
        row = table.rows[i].cells
        row[0].text = comp
        row[1].text = st
        row[2].text = det
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    return table


# ---------- documento ----------

def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---- Portada
    add_title(doc, "Constanza", size=30)
    add_subtitle(doc, "Plataforma de Cobranzas B2B Omnicanal")
    add_subtitle(doc, "Estado actual del proyecto — qué tenemos, cómo funciona y qué falta")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Fecha del informe: {date.today().strftime('%d/%m/%Y')}").italic = True
    doc.add_page_break()

    # ---- 1. Resumen ejecutivo
    add_h1(doc, "1. Resumen ejecutivo")
    add_paragraph(
        doc,
        "Constanza es una plataforma B2B para acelerar cobranzas y reducir mora. "
        "Combina gestión de cartera (clientes y facturas), comunicaciones omnicanal "
        "(Email, WhatsApp y voz), conciliación de pagos vía PSP (Cresium) e "
        "imputación manual con human-in-the-loop. A la fecha tiene un MVP "
        "desplegado en producción con los flujos principales operativos."
    )
    add_paragraph(doc, "Estado general: MVP funcional en producción.", bold=True)
    bullet(doc, "Frontend desplegado en Vercel y conectado al backend.")
    bullet(doc, "Backend (API + webhook Cresium + notifier) desplegado en Railway.")
    bullet(doc, "Base de datos Postgres y Redis productivos en Railway.")
    bullet(doc, "Flujo end-to-end de cobranza ya operativo: ingesta de facturas, "
                "envío de mensajes, recepción de pagos Cresium y conciliación.")

    # ---- 2. Stack y entornos
    add_h1(doc, "2. Stack tecnológico y entornos")
    add_h2(doc, "2.1 Stack")
    bullet(doc, "Frontend: Next.js 14 (App Router) + shadcn/ui + Tailwind, hosting en Vercel.")
    bullet(doc, "Backend: Node.js 20 + Fastify, hosting en Railway.")
    bullet(doc, "Base de datos: PostgreSQL (Railway/Supabase) con Prisma ORM.")
    bullet(doc, "Cache / colas: Redis (Railway) + BullMQ para procesamiento asíncrono.")
    bullet(doc, "Autenticación: JWT (HS256) entre web y API.")
    bullet(doc, "Mensajería: builderbot.cloud para WhatsApp, SMTP genérico (Nodemailer) para email.")
    bullet(doc, "Voz: ElevenLabs (TTS y agentes telefónicos) — integración parcial.")
    bullet(doc, "Pagos: Cresium (webhooks de depósitos / e-cheques / transferencias).")
    bullet(doc, "Logs estructurados con Pino.")

    add_h2(doc, "2.2 Entornos productivos actuales")
    bullet(doc, "Web (Vercel): frontend principal del producto.")
    bullet(doc, "api-gateway (Railway): API REST autenticada, expone toda la lógica de negocio a la UI.")
    bullet(doc, "rail-cucuru (Railway): webhook de Cresium para depósitos / e-cheques.")
    bullet(doc, "notifier (Railway): envío de WhatsApp y emails, polling/recepción y cola BullMQ.")
    bullet(doc, "Postgres (Railway): base de datos productiva.")
    bullet(doc, "Redis (Railway): colas e idempotencia.")

    # ---- 3. Módulos funcionales — qué tenemos hoy
    add_h1(doc, "3. Qué tenemos hoy (módulos funcionales)")

    add_h2(doc, "3.1 Autenticación y usuarios")
    bullet(doc, "Login con email/contraseña y emisión de JWT.")
    bullet(doc, "Perfiles: ADM, OPERADOR_1, OPERADOR_2 con permisos diferenciados.")
    bullet(doc, "Gestión de usuarios (pantalla en /settings/users).")

    add_h2(doc, "3.2 Clientes (cartera)")
    bullet(doc, "Alta y listado de clientes B2B con razón social, contactos y código único.")
    bullet(doc, "Carga masiva por archivo.")
    bullet(doc, "Búsqueda por teléfono normalizado (índice optimizado).")
    bullet(doc, "Soporte de varios CUIT por cliente.")

    add_h2(doc, "3.3 Facturas")
    bullet(doc, "Listado con estados (Abierta, Vencida, Parcial, Pagada).")
    bullet(doc, "Detalle de factura con timeline e historial de interacciones.")
    bullet(doc, "Carga manual y por archivo.")
    bullet(doc, "Cálculo de saldo aplicado a partir de imputaciones.")

    add_h2(doc, "3.4 Comunicaciones omnicanal")
    add_h3(doc, "WhatsApp (vía builderbot.cloud)")
    bullet(doc, "Envío de mensajes salientes (texto, plantillas).")
    bullet(doc, "Recepción de mensajes entrantes (texto, voz transcripta, imagen, documento).")
    bullet(doc, "Cobranza IA: respuesta automática con contexto de cliente y políticas por tenant.")
    bullet(doc, "Registro de todos los eventos en timeline de factura.")
    bullet(doc, "Extracción de promesas de pago y callbacks desde mensajes.")
    bullet(doc, "Pendiente solamente: definir y dar de alta el número de WhatsApp definitivo "
                "que va a usar el cliente.")
    add_h3(doc, "Email (SMTP)")
    bullet(doc, "Envío vía Nodemailer con cualquier proveedor SMTP (Gmail, Resend, SendGrid, etc.).")
    bullet(doc, "Recepción de respuestas vía webhook (procesado por OpenAI Agent en /wh/email/incoming).")
    bullet(doc, "Plantillas HTML responsivas para notificaciones.")
    add_h3(doc, "Voz / llamadas telefónicas (ElevenLabs ConvAI)")
    bullet(doc, "Integración completa con ElevenLabs Conversational AI: agente configurado, "
                "envío de llamadas salientes con variables dinámicas (cliente, factura, monto, etc.).")
    bullet(doc, "Carga de batches de llamadas desde Excel y ejecución masiva con rate limiting.")
    bullet(doc, "Listado de llamadas, detalle, duración y resumen automático.")
    bullet(doc, "Cronograma de callbacks (recordatorios programados) y reintentos.")
    bullet(doc, "Extracción automática de promesas y callbacks desde el resumen de la llamada.")
    bullet(doc, "Registro de cada llamada en el timeline de la factura.")
    bullet(doc, "TTS de ElevenLabs disponible también para mensajes de voz en WhatsApp.")
    bullet(doc, "Pendiente solamente: cargar el número físico (ELEVENLABS_PHONE_NUMBER_ID) "
                "que va a usar el cliente en producción.")

    add_h2(doc, "3.5 Pagos y conciliación")
    bullet(doc, "Webhook firmado de Cresium en /wh/cresium/deposito con verificación HMAC.")
    bullet(doc, "Persistencia en pay.payments + pay.payment_applications.")
    bullet(doc, "Matching automático por número de factura, CUIT o CVU del ordenante.")
    bullet(doc, "Diferenciación entre transferencia bancaria (TRANSFERENCIA) y e-cheque (ECHEQ).")
    bullet(doc, "Pantalla \"Transferencias bancarias\" con datos del pagador (CVU, CUIT, nombre).")
    bullet(doc, "Pantalla \"E-cheques\" dedicada con consulta a API Cresium v3 (Partner).")
    bullet(doc, "Pantalla \"Conciliación de Pagos\" para imputar manualmente y aprobar/rechazar.")
    bullet(doc, "Aprobar / Liquidar pago desde el detalle del pago (acción rápida).")
    bullet(doc, "Notificación automática por WhatsApp al recibir un cobro (empresa y cliente).")
    bullet(doc, "Idempotencia con Redis y deduplicación por externalRef.")

    add_h2(doc, "3.6 Cobranza inteligente")
    bullet(doc, "Políticas por tenant configurables.")
    bullet(doc, "Prompt dinámico construido con contexto de cliente y deuda.")
    bullet(doc, "Llamada a OpenAI para responder automáticamente al deudor en WhatsApp.")
    bullet(doc, "Persistencia de respuesta y tokens consumidos en el evento.")

    add_h2(doc, "3.7 KPIs y dashboards")
    bullet(doc, "Dashboard con KPI principales (cobrado, pendiente, e-cheques pend. liquidación, etc.).")
    bullet(doc, "Pantalla de jobs internos para tareas programadas.")
    bullet(doc, "Listado de mensajes enviados y su progreso.")

    add_h2(doc, "3.8 Integraciones administrativas")
    bullet(doc, "Proxy interno a API Partner de Cresium v3 (consulta de transacciones).")
    bullet(doc, "Variables de entorno y RLS controladas (con desactivación temporal en tablas de negocio "
                "para estabilidad detrás de pooler).")

    # ---- 4. Estado por componente
    doc.add_page_break()
    add_h1(doc, "4. Estado por componente")
    rows = [
        ("Frontend (apps/web)", "OK — productivo",
         "Login, dashboard, clientes, facturas, pagos (transferencias, e-cheques, conciliación), "
         "notificaciones, llamadas, jobs y administración de usuarios."),
        ("API Gateway (apps/api-gateway)", "OK — productivo",
         "Rutas: auth, customers, invoices, payments, kpi, notify, calls, jobs, integrations, "
         "tenants, tenant-settings, users, cresium-partner, agent-context, summaries, chat, seed."),
        ("Webhook Cresium (apps/rail-cucuru)", "OK — productivo",
         "Recibe depósitos firmados, identifica método (transferencia / e-cheque), aplica matching "
         "automático y notifica vía notifier."),
        ("Notifier (apps/notifier)", "OK — productivo",
         "Envío de WhatsApp/email, polling y recepción, cola BullMQ, IA de cobranza."),
        ("Postgres (Railway)", "OK — productivo",
         "Esquemas core, pay, contact, ops y bindx. RLS desactivada temporalmente en algunas tablas "
         "por compatibilidad con el pooler — pendiente revisión."),
        ("Redis (Railway)", "OK — productivo",
         "Cola BullMQ e idempotencia de webhooks."),
        ("Integración Cresium API v3 (Partner)", "OK — funcional",
         "Proxy seguro desde api-gateway con firma HMAC; pantalla de detalle de transacción."),
        ("ElevenLabs — voz / llamadas", "OK funcional — falta número físico",
         "Integración completa con ElevenLabs ConvAI: envío de llamadas, variables dinámicas, "
         "batches, listado, cronograma de callbacks y resumen con extracción de promesas. "
         "Pendiente: aprovisionar el número saliente definitivo (ELEVENLABS_PHONE_NUMBER_ID)."),
        ("WhatsApp (BuilderBot)", "OK funcional — falta número definitivo",
         "Envío y recepción operativos, cobranza IA funcionando. Pendiente: dar de alta el "
         "número de WhatsApp definitivo del cliente."),
        ("BindX — e-cheques (rail-bindx)", "Pendiente",
         "Cresium ya cubre e-cheques operativos; BindX queda en roadmap si se decide ampliar."),
        ("Reconciler dedicado", "No requerido aún",
         "El matching automático ya vive dentro de rail-cucuru; un servicio reconciler separado "
         "queda para escala mayor."),
        ("Doc-intake / OCR", "Pendiente",
         "Recepción de comprobantes pero sin pipeline de OCR/validación automática todavía."),
        ("ERP Connector", "Pendiente",
         "Export hacia ERP del cliente: aún no implementado."),
        ("Observabilidad avanzada", "Básico",
         "Logs estructurados con Pino. Métricas/Alertas (OpenTelemetry) no implementadas."),
    ]
    add_status_table(doc, rows)

    # ---- 5. Flujos end-to-end probados
    add_h1(doc, "5. Flujos end-to-end ya operativos")
    add_h2(doc, "5.1 Cobranza por WhatsApp con IA")
    bullet(doc, "Cliente responde un WhatsApp.")
    bullet(doc, "Notifier recibe el evento, lo asocia a cliente y factura.")
    bullet(doc, "IA construye respuesta con contexto y políticas del tenant.")
    bullet(doc, "Se envía la respuesta y se registra todo en el timeline.")

    add_h2(doc, "5.2 Cobro por transferencia / e-cheque (Cresium)")
    bullet(doc, "Cresium dispara webhook firmado a rail-cucuru.")
    bullet(doc, "Se valida firma HMAC, ventana anti-replay y monto.")
    bullet(doc, "Se intenta match automático por número de factura, CUIT o CVU.")
    bullet(doc, "Si concilia → se crea payment_applications autoritativo + se liquida.")
    bullet(doc, "Si no concilia → queda en \"Conciliación de Pagos\" para imputación manual.")
    bullet(doc, "Se notifica por WhatsApp a la empresa y al cliente identificado.")

    add_h2(doc, "5.3 Aprobación humana desde frontend")
    bullet(doc, "Operador ve pagos pendientes en /payments/reconciliation.")
    bullet(doc, "Imputa el pago a una factura sugerida (candidatos automáticos).")
    bullet(doc, "Aprueba (LIQUIDATE) o rechaza (REJECT) desde la pantalla.")
    bullet(doc, "Desde el detalle del pago también puede aprobar en un clic.")

    # ---- 6. Cambios y mejoras más recientes
    add_h1(doc, "6. Cambios recientes destacados")
    bullet(doc, "Aprobar y liquidar pagos desde el detalle del pago (frontend).")
    bullet(doc, "Módulo de e-cheques Cresium: ruta, KPI, proxy a API Partner y metadata de transacción.")
    bullet(doc, "Notificación automática de cobros por WhatsApp a empresa y cliente.")
    bullet(doc, "Robustecimiento de envío a BuilderBot con fallbacks y errores HTTP detallados.")
    bullet(doc, "Estabilización del backend (RLS, queries de clientes y facturas).")
    bullet(doc, "Normalización de URL de NOTIFIER_URL y firma de Cresium tolerante a variantes.")

    # ---- 7. Qué falta para llegar al 100%
    doc.add_page_break()
    add_h1(doc, "7. Qué falta para dejar lo actual 100% funcional")

    add_h2(doc, "7.1 Bloqueantes / críticos")
    bullet(doc, "Cresium en producción con firma activa: revisar que en rail-cucuru "
                "CRESIUM_SKIP_SIGNATURE_VERIFY esté en false y el secret correcto cargado.")
    bullet(doc, "Reactivar RLS o equivalente en las tablas de negocio que quedaron temporalmente sin RLS, "
                "garantizando aislamiento por tenant a nivel base de datos.")
    bullet(doc, "Auditoría de variables de entorno en Railway/Vercel: validar que cada servicio "
                "tenga las claves productivas correctas (Cresium, SMTP, BuilderBot, ElevenLabs, JWT).")

    add_h2(doc, "7.2 Aprovisionamiento de canales (no es código — es alta)")
    bullet(doc, "WhatsApp: definir y dar de alta el número definitivo del cliente en BuilderBot Cloud "
                "y cargar credenciales en notifier. El resto del flujo ya está operativo.")
    bullet(doc, "Voz / llamadas: contratar y cargar el número saliente en ElevenLabs "
                "(ELEVENLABS_PHONE_NUMBER_ID). La integración con ConvAI ya está hecha y probada.")
    bullet(doc, "Email: validar dominio del cliente con el proveedor SMTP elegido (Resend/SendGrid/etc.).")

    add_h2(doc, "7.3 Cierre funcional del alcance comprometido")
    bullet(doc, "Conciliación: caso de rechazo con motivo y notificación al área operativa.")
    bullet(doc, "Reintentos en cola para mensajes fallidos (BullMQ) con visibilidad en UI.")
    bullet(doc, "Acción rápida de \"Rechazar\" en el detalle del pago (hoy queda en conciliación).")
    bullet(doc, "Cierre del módulo de jobs programados (T-3 / T-1 / D0) con scheduler real.")
    bullet(doc, "Reportes / exportaciones (CSV) de pagos y facturas.")

    add_h2(doc, "7.4 Calidad y operación")
    bullet(doc, "Tests automatizados para los endpoints críticos (auth, pagos, webhooks).")
    bullet(doc, "Monitoreo y alertas (errores 5xx, fallas SMTP, fallas firma Cresium, colas atascadas).")
    bullet(doc, "Backups verificados de Postgres con prueba de restore.")
    bullet(doc, "Documentación operativa para el equipo (runbooks).")
    bullet(doc, "Procedimiento de rollback documentado por servicio.")

    add_h2(doc, "7.5 Producto / UX")
    bullet(doc, "Estandarizar mensajes de error en la UI (cobranza, pagos, carga de archivos).")
    bullet(doc, "Filtros y exportación en listados (clientes, facturas, pagos).")
    bullet(doc, "Pantalla unificada del estado del cliente (deuda + interacciones + pagos).")

    add_h2(doc, "7.6 Seguridad")
    bullet(doc, "Rotación periódica de JWT_SECRET y secrets de webhooks.")
    bullet(doc, "Revisar políticas de rate limit en api-gateway.")
    bullet(doc, "Revisión de logs para que no expongan datos sensibles (CVU, CUIT, etc.).")

    # ---- 8. Roadmap fuera del alcance actual
    add_h1(doc, "8. Roadmap fuera del alcance \"100% actual\"")
    bullet(doc, "rail-bindx (BindX e-cheques).")
    bullet(doc, "reconciler como microservicio dedicado.")
    bullet(doc, "doc-intake con OCR de comprobantes.")
    bullet(doc, "erp-connector y export periódico al ERP del cliente.")
    bullet(doc, "commissions-service para asociados.")
    bullet(doc, "Métricas / observabilidad con OpenTelemetry.")
    bullet(doc, "Portal del cliente deudor (autoatención).")

    # ---- 9. Prueba de circuito completo (end-to-end)
    doc.add_page_break()
    add_h1(doc, "9. Prueba de circuito completo (end-to-end)")
    add_paragraph(
        doc,
        "Objetivo: ejecutar un caso real de punta a punta para validar todo el flujo "
        "de cobranza: cargar una factura, gestionar la cobranza por los tres canales "
        "(email, WhatsApp y llamada), recibir el pago vía Cresium y dejarlo conciliado."
    )

    add_h2(doc, "9.1 Pasos del circuito a probar")
    bullet(doc, "1) Crear / dar de alta el cliente deudor con email, teléfono y CUIT.")
    bullet(doc, "2) Cargar una factura abierta a ese cliente (manual o importación).")
    bullet(doc, "3) Ejecutar gestión de cobranza:")
    bullet(doc, "   - Envío de email recordatorio (plantilla productiva).", level=1)
    bullet(doc, "   - Envío de WhatsApp con plantilla y prueba de IA respondiendo.", level=1)
    bullet(doc, "   - Llamada saliente con ElevenLabs ConvAI y resumen automático.", level=1)
    bullet(doc, "4) Confirmar registro de cada interacción en el timeline de la factura.")
    bullet(doc, "5) Simular el pago: el deudor transfiere o emite e-cheque vía Cresium.")
    bullet(doc, "6) Verificar que Cresium dispara el webhook a rail-cucuru.")
    bullet(doc, "7) Verificar matching automático contra la factura (o caer a Conciliación).")
    bullet(doc, "8) Aprobar / liquidar el pago desde el frontend.")
    bullet(doc, "9) Verificar que la factura queda SALDADA y que se notifica por WhatsApp.")

    add_h2(doc, "9.2 Qué falta concretamente para correr esta prueba")
    add_h3(doc, "Aprovisionamiento (no requiere desarrollo)")
    bullet(doc, "Número de WhatsApp definitivo dado de alta en BuilderBot Cloud "
                "con sus credenciales cargadas en Railway (notifier).")
    bullet(doc, "Número saliente de voz aprovisionado y cargado en ElevenLabs "
                "(ELEVENLABS_PHONE_NUMBER_ID + agente productivo).")
    bullet(doc, "Casilla / dominio de email validado en el proveedor SMTP "
                "(Resend / SendGrid / Gmail corporativo, etc.).")
    bullet(doc, "Cuenta Cresium en modo productivo con el secret correcto y el CVU "
                "de cobro configurado en el tenant.")

    add_h3(doc, "Datos de prueba")
    bullet(doc, "Tenant configurado con CVU de cobro real (en /payments/reconciliation).")
    bullet(doc, "Un cliente real (o de prueba) con email + WhatsApp + CUIT cargados.")
    bullet(doc, "Una factura cargada con monto exacto al que se va a transferir, "
                "para validar el matching automático.")

    add_h3(doc, "Validaciones técnicas previas")
    bullet(doc, "Verificar que CRESIUM_SKIP_SIGNATURE_VERIFY esté en false en producción.")
    bullet(doc, "Verificar que el webhook de Cresium responde 200 a /wh/cresium/deposito.")
    bullet(doc, "Verificar envío real de email (no rechazo por SPF/DKIM).")
    bullet(doc, "Verificar que WhatsApp entrega el mensaje al número del cliente.")
    bullet(doc, "Verificar que ElevenLabs marca correctamente el outbound.")

    add_h3(doc, "Cierre operativo de la prueba")
    bullet(doc, "Definir el caso de éxito esperado (factura saldada y notificaciones enviadas).")
    bullet(doc, "Definir casos negativos a probar (sin match → cae en Conciliación; "
                "monto distinto; rechazo manual).")
    bullet(doc, "Designar 1 operador para correr la prueba y 1 técnico para monitorear logs.")

    # ---- 10. Integración con ERPs (API manager)
    add_h1(doc, "10. Integración con ERPs (capa de API manager)")
    add_paragraph(
        doc,
        "Una vez validado el circuito interno, el siguiente bloque grande es "
        "conectar Constanza con los ERPs de cada cliente. Esto requiere una capa "
        "de integración tipo API manager / connector que estandarice cómo se "
        "envía y se toma información, independientemente del ERP de turno "
        "(SAP, Oracle, Tango, Bejerman, Calipso, Odoo, Holistor, etc.)."
    )

    add_h2(doc, "10.1 Qué tiene que hacer esa capa")
    bullet(doc, "Recibir facturas desde el ERP (push o pull) y crear/actualizar en Constanza.")
    bullet(doc, "Enviar al ERP los pagos imputados / cobros liquidados (export).")
    bullet(doc, "Mantener equivalencias entre IDs internos de Constanza e IDs del ERP.")
    bullet(doc, "Reintentos, idempotencia y registro de errores de sincronización.")
    bullet(doc, "Mapeo configurable de campos por cliente (cada ERP tiene su esquema).")
    bullet(doc, "Auditoría: qué se envió, cuándo, con qué resultado.")

    add_h2(doc, "10.2 Opciones de arquitectura")
    bullet(doc, "Microservicio propio (erp-connector) ya previsto en el roadmap del "
                "monorepo, con un adaptador por ERP.")
    bullet(doc, "API manager externo (n8n / Make / Zapier / Workato / Mulesoft) "
                "orquestando llamadas entre el ERP y la API de Constanza.")
    bullet(doc, "Híbrido: un erp-connector básico en Constanza + un orquestador "
                "externo para los casos más exóticos.")

    add_h2(doc, "10.3 Endpoints que Constanza tiene que exponer (alto nivel)")
    bullet(doc, "POST /ingest/invoices — alta o upsert de facturas desde el ERP.")
    bullet(doc, "POST /ingest/customers — alta o upsert de clientes desde el ERP.")
    bullet(doc, "GET /exports/payments — pagos imputados en un rango de fechas.")
    bullet(doc, "GET /exports/invoices — estado actual de facturas por cliente.")
    bullet(doc, "Webhook saliente \"payment.applied / settled\" para que el ERP "
                "se entere en tiempo real.")

    add_h2(doc, "10.4 Qué falta para arrancar esta etapa")
    bullet(doc, "Definir con cada cliente cuál es su ERP y qué prefiere: push, pull "
                "o usar un API manager externo.")
    bullet(doc, "Construir el microservicio erp-connector (o configurar el flujo en n8n/Make).")
    bullet(doc, "Definir el contrato de datos (mapeo de campos por ERP).")
    bullet(doc, "Mecanismo de autenticación entre ERP y Constanza (API keys por cliente).")
    bullet(doc, "Sandbox de pruebas con datos del ERP real para no romper producción.")
    bullet(doc, "Monitoreo de la sincronización (qué falla, qué se reintenta, qué quedó pendiente).")

    # ---- 11. Conclusión
    add_h1(doc, "11. Conclusión")
    add_paragraph(
        doc,
        "El MVP de Constanza está vivo y resolviendo el flujo principal del negocio: "
        "comunicación con el deudor por email, WhatsApp y voz, recepción de pagos "
        "vía Cresium, conciliación (automática y manual) y aprobación humana desde "
        "la web. El esfuerzo remanente para llegar al 100% del alcance comprometido "
        "se concentra en aprovisionar los canales (número de WhatsApp y número de "
        "voz), endurecer la plataforma (RLS, firmas, monitoreo) y cerrar detalles "
        "operativos (rechazos, reintentos, reportes)."
    )
    add_paragraph(
        doc,
        "El siguiente gran bloque es la prueba real end-to-end (factura → cobranza "
        "omnicanal → pago Cresium → conciliación) y, en paralelo, la conexión con "
        "los ERPs de los clientes a través de una capa de integración tipo API "
        "manager. Con esos dos pasos cerrados, Constanza queda listo para escalar "
        "a más tenants y sumar los módulos del roadmap."
    )

    out = "Constanza-Estado-del-Proyecto.docx"
    doc.save(out)
    print(f"Documento generado: {out}")


if __name__ == "__main__":
    main()
