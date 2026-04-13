#!/usr/bin/env python3
"""Genera documento Word detallado: Constanza — funcionalidades completas."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def p(doc, text: str, bold: bool = False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para


def bullet(doc, text: str):
    doc.add_paragraph(text, style="List Bullet")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Constanza")
    r.bold = True
    r.font.size = Pt(24)
    doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run(
        "Descripción funcional integral — Cobranzas, contact center, conciliación e integraciones"
    ).italic = True
    doc.add_paragraph(
        "Documento para negocio, operación y TI. Incluye visión de producto, dominios de datos, "
        "flujos principales y estado de implementación según el monorepo actual."
    )
    doc.add_page_break()

    # --- 1
    doc.add_heading("1. Resumen ejecutivo", 1)
    doc.add_paragraph(
        "Constanza es una plataforma B2B de cobranzas y gestión de cartera que combina: "
        "(1) comunicación omnicanal con deudores (email, WhatsApp, voz), "
        "(2) cartera de clientes y facturas con estados y saldos, "
        "(3) recepción e imputación de pagos (transferencias vía PSP como Cresium, imputación manual u otros orígenes), "
        "(4) políticas de cobranza y seguimiento (promesas, callbacks), "
        "(5) reporting y, donde aplica, intervención humana."
    )

    # --- 2
    doc.add_heading("2. Visión del producto", 1)
    doc.add_paragraph(
        "El objetivo operativo es reducir mora y tiempo de cobro con trazabilidad end-to-end: "
        "desde el primer contacto hasta el pago imputado y el cierre contable en la factura."
    )
    bullet(doc, "Unificar en un solo lugar la conversación con el cliente y el estado real de la deuda.")
    bullet(doc, "Automatizar conciliación cuando el PSP o el banco informa el cobro (webhooks firmados).")
    bullet(doc, "Mantener políticas de tono, horarios y límites de negociación por empresa (tenant).")

    # --- 3
    doc.add_heading("3. Arquitectura técnica (alto nivel)", 1)
    doc.add_paragraph(
        "Stack principal: frontend Next.js (Vercel), API Fastify (Railway), PostgreSQL multi-schema "
        "(Supabase/Railway), Redis para colas (BullMQ), autenticación JWT y RLS por tenant."
    )
    doc.add_heading("3.1 Servicios y responsabilidades", 2)
    doc.add_paragraph(
        "En el monorepo coexisten aplicaciones con distinto grado de madurez. "
        "Los núcleos operativos desplegados típicamente incluyen: web, api-gateway, notifier y rail-cucuru (Cresium)."
    )
    bullet(doc, "apps/web: interfaz de usuario (facturas, clientes, transferencias, notificaciones, llamadas según pantallas).")
    bullet(doc, "apps/api-gateway: API REST autenticada (JWT), reglas de negocio expuestas a la UI y a agentes externos (API key).")
    bullet(doc, "apps/notifier: envío de email (SMTP), WhatsApp (BuilderBot Cloud), cola BullMQ, registro en contact.events.")
    bullet(doc, "apps/rail-cucuru: único webhook de ingresos Cresium (depósitos/transferencias), validación HMAC, persistencia en pay.payments.")
    doc.add_paragraph(
        "Otros microservicios descritos en la arquitectura (reconciler dedicado, contact-orchestrator standalone, etc.) "
        "pueden estar en roadmap o parcialmente integrados; el modelo de datos ya los soporta."
    )

    # --- 4
    doc.add_heading("4. Multi-tenant, seguridad y datos", 1)
    doc.add_paragraph(
        "Cada empresa es un tenant. Los datos de negocio están aislados lógicamente; en PostgreSQL suele aplicarse "
        "Row Level Security (RLS) con contexto de sesión (tenant_id) en las consultas."
    )
    bullet(doc, "Usuarios internos: perfiles ADM, OPERADOR_1, OPERADOR_2 con permisos diferenciados.")
    bullet(doc, "Clientes deudores pueden existir como registros en core.customers; el acceso portal cliente es opcional según diseño.")
    bullet(doc, "Webhooks externos (Cresium) se autentican por firma HMAC y variables de entorno; los agentes de IA usan AGENT_API_KEY en api-gateway.")

    # --- 5
    doc.add_heading("5. Dominio core: cartera y facturación", 1)
    doc.add_heading("5.1 Clientes (core.customers)", 2)
    doc.add_paragraph(
        "Representa al deudor B2B: razón social, email, teléfono (clave para WhatsApp y matching), "
        "código único de negocio (p. ej. CVU esperado en transferencias para conciliar), referencias externas al ERP."
    )
    doc.add_heading("5.2 CUITs (core.customer_cuits)", 2)
    doc.add_paragraph(
        "Un cliente puede tener varios CUIT; se usa para cruzar avisos de pago del PSP con la cartera."
    )
    doc.add_heading("5.3 Facturas (core.invoices)", 2)
    doc.add_paragraph(
        "Factura con número, monto en centavos, fecha de vencimiento y estado (ABIERTA, VENCIDA, PARCIAL, PAGADA/SALDADA, etc. según reglas de negocio). "
        "Es la unidad principal de cobranza y de imputación de pagos."
    )
    doc.add_heading("5.4 Promesas de pago (core.promises)", 2)
    doc.add_paragraph(
        "Compromisos registrados por canal (email, WhatsApp, voz): monto opcional, fecha de cumplimiento, estado (pendiente/cumplida/incumplida) y motivo."
    )
    doc.add_heading("5.5 Políticas (core.policy_rules)", 2)
    doc.add_paragraph(
        "Parámetros por tenant en JSON: tono, horarios de contacto, máximo de negociaciones, porcentaje mínimo de pago, etc. "
        "Alimentan prompts de IA y reglas de contacto."
    )

    # --- 6 Contact center
    doc.add_heading("6. Contact center y omnicanalidad", 1)
    doc.add_paragraph(
        "Constanza modela el contacto como eventos auditables enlazados a cliente y, cuando aplica, a factura."
    )
    doc.add_heading("6.1 Secuencias y ejecución (contact.sequences, contact.runs)", 2)
    doc.add_paragraph(
        "Las secuencias definen pasos (email, WhatsApp, voz) con retardos relativos al vencimiento, ventanas horarias y límites de intentos. "
        "Cada run asocia una factura (y opcionalmente cliente) a un estado de campaña (próximo a vencer, vencida, promesa, etc.)."
    )
    doc.add_heading("6.2 Eventos de contacto (contact.events)", 2)
    doc.add_paragraph(
        "Cada envío o recepción se registra con canal, dirección (OUTBOUND/INBOUND), texto, estado, IDs externos del proveedor, "
        "posible URL de medio, transcripción o resumen de llamada, y payload JSON para contexto adicional."
    )
    doc.add_heading("6.3 Lotes y operación masiva (contact.batch_jobs)", 2)
    doc.add_paragraph(
        "Envíos masivos por canal con seguimiento de procesados y fallidos; integrado con notifier y colas."
    )
    doc.add_heading("6.4 Callbacks programados (contact.scheduled_callbacks)", 2)
    doc.add_paragraph(
        "Seguimientos con fecha/hora, tipo (CALLBACK o FOLLOW_UP), motivo y vínculo opcional al evento que los originó. "
        "Permite agenda operativa y cumplimiento de compromisos conversacionales."
    )
    doc.add_heading("6.5 Canales", 2)
    bullet(doc, "Email: SMTP (Gmail, SendGrid, etc.); asunto y cuerpo con variables.")
    bullet(doc, "WhatsApp: integración BuilderBot Cloud; texto, notas de voz, imagen, documento según configuración.")
    bullet(doc, "Voz: modelo unificado en eventos; integración con agentes de voz/TTS según despliegue.")

    # --- 7 IA
    doc.add_heading("7. Inteligencia artificial y contexto para agentes", 1)
    doc.add_paragraph(
        "El notifier puede enriquecer respuestas con políticas de cobranza y contexto de cliente/factura desde base de datos. "
        "Los mensajes entrantes de WhatsApp pueden disparar flujos que generan respuesta y la registran en el timeline."
    )
    doc.add_heading("7.1 API de contexto para agentes (api-gateway)", 2)
    doc.add_paragraph(
        "GET /v1/agent/context autenticado con AGENT_API_KEY: devuelve cliente, facturas según scope (open=pending, paid=pagadas, all=historial amplio), "
        "totales y resumen textual para el bot."
    )
    doc.add_heading("7.2 Acciones desde el bot", 2)
    doc.add_paragraph(
        "POST /v1/agent/actions/promise: registra promesa de pago vinculada a factura. "
        "POST /v1/agent/actions/callback: agenda seguimiento en scheduled_callbacks. "
        "Requieren identificar cliente por customerId, phone o email en el cuerpo JSON."
    )

    # --- 8 Pagos
    doc.add_heading("8. Pagos, imputación y conciliación", 1)
    doc.add_heading("8.1 Modelo pay.payments y pay.payment_applications", 2)
    doc.add_paragraph(
        "Un pago (origen Cresium, manual u otro) se almacena con método, estado, referencia externa, metadata JSON (payload PSP, CUIT/CVU extraídos) "
        "y opcionalmente monto total si aún no hay aplicaciones. Las aplicaciones vinculan pago a factura con monto en centavos y marca is_authoritative cuando el origen es definitivo."
    )
    doc.add_heading("8.2 Integración Cresium (apps/rail-cucuru)", 2)
    doc.add_paragraph(
        "El servicio expone POST de depósito; valida firma HMAC (CRESIUM_PARTNER_SECRET / CRESIUM_WEBHOOK_SECRET), "
        "anti-replay por timestamp, idempotencia con Redis y external_ref. "
        "Intenta imputar automáticamente comparando número de factura en texto, CUIT/CVU del cliente en cartera, "
        "y reglas opcionales (p. ej. match solo por monto si está habilitado explícitamente). "
        "Variables clave: CRESIUM_TENANT_ID, secretos de firma, opcionalmente CRESIUM_REJECT_CVU_MISMATCH, CRESIUM_AUTO_MATCH_AMOUNT_ONLY, etc."
    )
    doc.add_heading("8.3 Conciliación y reglas de negocio", 2)
    doc.add_paragraph(
        "La visión de arquitectura incluye un reconciler para casos ambiguos (match parcial, 1 a N, tolerancias) y cola de decisiones para excepciones humanas. "
        "Las aplicaciones con is_authoritative=true no deben ser sobrescritas silenciosamente por reglas automáticas."
    )
    doc.add_heading("8.4 Notificaciones post-pago (operativo)", 2)
    doc.add_paragraph(
        "Tras registrar un depósito puede enviarse WhatsApp al receptor interno y confirmación al pagador si hay teléfono e imputación, "
        "vía notifier y variables de entorno en rail-cucuru (NOTIFIER_URL, etc.)."
    )

    # --- 9 Banco / PSP
    doc.add_heading("9. Conexión con banco y proveedor de pagos (PSP)", 1)
    doc.add_paragraph(
        "La “conexión con el banco” en la práctica se materializa como integración con el proveedor de cobros (en este despliegue, Cresium): "
        "webhooks de acreditación/depósito, validación criptográfica, y alineación del tenant destino (CRESIUM_TENANT_ID) con la empresa cobradora en Constanza."
    )
    bullet(doc, "CVU de cobro de la empresa puede almacenarse en tenant para validación estricta opcional.")
    bullet(doc, "Los avisos incluyen montos y a veces identificadores fiscales o CVU para cruzar con clientes.")

    # --- 10 Otros
    doc.add_heading("10. Otros dominios y extensiones", 1)
    doc.add_heading("10.1 E-cheques / BindX (visión)", 2)
    doc.add_paragraph(
        "Arquitectura prevista: listado y aceptación/rechazo de e-cheques, webhooks de liquidación y creación de aplicaciones de pago."
    )
    doc.add_heading("10.2 Asociados y comisiones", 2)
    doc.add_paragraph(
        "Tabla core.asociados para futuros cálculos de comisión por vendedor/asociado ligados a cobranzas."
    )
    doc.add_heading("10.3 KPI y reporting", 2)
    doc.add_paragraph(
        "Endpoints y pantallas de KPI según módulo web; datos agregados desde facturas, pagos y eventos de contacto."
    )

    # --- 11 API
    doc.add_heading("11. Superficie API principal (api-gateway)", 1)
    doc.add_paragraph(
        "Prefijo típico /v1 salvo /auth y /health. Autenticación JWT para usuarios; rutas /v1/agent/* con AGENT_API_KEY."
    )
    bullet(doc, "/auth: login y sesión.")
    bullet(doc, "/v1/invoices, /v1/customers: alta y consulta de facturas y clientes.")
    bullet(doc, "/v1/payments: pagos e imputaciones según reglas del servicio.")
    bullet(doc, "/v1/kpi: indicadores agregados.")
    bullet(doc, "/v1/notify: disparo de notificaciones desde la plataforma.")
    bullet(doc, "/v1/calls: gestión de llamadas y timeline asociado.")
    bullet(doc, "/v1/tenant-settings, /v1/integrations: configuración del tenant e integraciones.")
    bullet(doc, "/v1/summaries, /v1/jobs: resúmenes y trabajos en segundo plano.")
    bullet(doc, "/v1/cobranza: políticas y datos de cobranza.")
    bullet(doc, "/v1/users, /v1/tenants: usuarios internos y listado de tenants (según permisos).")
    bullet(doc, "/v1/agent/context, /v1/agent/actions/promise, /v1/agent/actions/callback: contexto y acciones para bots.")

    # --- 12 Cierre
    doc.add_heading("12. Conclusión", 1)
    doc.add_paragraph(
        "Constanza integra gestión de cobranzas, contact center omnicanal, políticas configurables, "
        "recepción de pagos por PSP con imputación y trazabilidad, y APIs para agentes conversacionales. "
        "La profundidad de cada módulo depende del despliegue y roadmap; este documento refleja el diseño objetivo y los componentes presentes en el repositorio."
    )

    out = "/Users/ralborta/Constanza/docs/Constanza-Descripcion-Funcional-Completa.docx"
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
