#!/usr/bin/env python3
"""Genera documento Word con descripción funcional de Constanza."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Constanza")
    r.bold = True
    r.font.size = Pt(22)
    doc.add_paragraph().add_run(
        "Sistema de cobranzas B2B omnicanal — descripción funcional"
    ).italic = True
    doc.add_paragraph(
        "Documento orientado a negocio y operación. "
        "Versión alineada con la arquitectura del monorepo y los servicios desplegados."
    )
    doc.add_paragraph()

    # 1
    add_heading(doc, "1. Objetivo del producto", 1)
    doc.add_paragraph(
        "Constanza acelera y ordena la cobranza B2B combinando: "
        "comunicación multicanal con el deudor, registro unificado de interacciones, "
        "políticas de cobranza parametrizables, imputación y conciliación de pagos "
        "(transferencias y otros medios), reporting y, cuando hace falta, intervención humana "
        "(cola de decisiones)."
    )

    # 2
    add_heading(doc, "2. Arquitectura lógica (visión de alto nivel)", 1)
    doc.add_paragraph(
        "El ecosistema se organiza en microservicios con dominios claros: "
        "front web (Next.js), API de acceso (api-gateway), notificaciones (notifier), "
        "recepción de pagos por webhooks (rail-cucuru para Cresium, entre otros), "
        "colas (Redis + BullMQ) y base de datos PostgreSQL multi-tenant con RLS."
    )
    doc.add_paragraph(
        "Principios: multi-tenant, aislamiento por tenant, idempotencia en operaciones críticas "
        "(pagos, webhooks) y trazabilidad mediante eventos en contacto y en pagos."
    )

    # 3 Contact center
    add_heading(doc, "3. Capacidad de contact center", 1)
    doc.add_paragraph(
        "Constanza actúa como núcleo de contacto con el cliente deudor enlazado a la factura: "
        "cada interacción puede quedar asociada a cliente y, cuando aplica, a factura concreta."
    )
    add_heading(doc, "3.1 Modelo de datos de contacto", 2)
    doc.add_paragraph(
        "Esquema contact: secuencias (sequences), ejecuciones de secuencia por factura (runs) "
        "y eventos (events) que registran salidas (OUTBOUND) y entradas (INBOUND) por canal."
    )
    doc.add_paragraph(
        "Los eventos permiten armar el historial operativo: qué se envió, cuándo, por qué canal, "
        "estado (enviado, entregado, fallido) y metadatos para auditoría."
    )
    add_heading(doc, "3.2 Secuencias y orquestación", 2)
    doc.add_paragraph(
        "Las secuencias se definen en JSON: pasos con canal (EMAIL, WHATSAPP, VOICE), "
        "plantillas o mensajes, retardos relativos al vencimiento (por ejemplo T-3, D0), "
        "ventanas horarias y límites de intentos por canal. "
        "La visión de producto incluye cooldown por cliente para no saturar."
    )
    doc.add_paragraph(
        "El contact orchestrator (componente previsto en arquitectura) es quien, en producción completa, "
        "dispararía cohortes (próximo a vencer, vencida, promesa incumplida) según reglas de negocio."
    )

    # 4 Contexto
    add_heading(doc, "4. Manejo de contextos e “historia clínica”", 1)
    doc.add_paragraph(
        "“Contexto” en Constanza no es solo el mensaje actual: se construye a partir de "
        "datos maestros (cliente, facturas, saldos, estado de factura) y del timeline de "
        "contact.events asociado al cliente y/o a la factura."
    )
    add_heading(doc, "4.1 Contexto para agentes de IA (cobranza)", 2)
    doc.add_paragraph(
        "El notifier puede enriquecer respuestas usando políticas por tenant (policy_rules) "
        "y contexto de cliente/factura obtenido desde la base (por ejemplo tono, horarios, "
        "máximo de negociaciones, medios de pago aceptados). "
        "Eso alimenta prompts dinámicos y respuestas coherentes con la política comercial."
    )
    add_heading(doc, "4.2 Historial por factura y por cliente", 2)
    doc.add_paragraph(
        "En la interfaz, el historial de una factura agrega: mensajes salientes y entrantes, "
        "llamadas, promesas y callbacks derivados del análisis del texto. "
        "Así el operador ve la “historia clínica” del caso: qué se prometió, qué respondió el cliente "
        "y qué acciones automáticas se generaron."
    )

    # 5 Omnicanalidad
    add_heading(doc, "5. Omnicanalidad: Email, WhatsApp y voz", 1)
    add_heading(doc, "5.1 Email", 2)
    doc.add_paragraph(
        "Canal EMAIL: envío vía SMTP (proveedores típicos: Gmail con app password, SendGrid, etc.). "
        "Soporta asunto y cuerpo con variables. Los envíos masivos pueden agruparse en batch jobs "
        "con seguimiento de procesados y fallidos."
    )
    add_heading(doc, "5.2 WhatsApp", 2)
    doc.add_paragraph(
        "Canal WHATSAPP: integración con BuilderBot Cloud (API de envío). "
        "Soporta texto, notas de voz, imagen y documento según configuración del proveedor. "
        "Los mensajes entrantes pueden disparar flujo de IA de cobranza y registro en timeline."
    )
    add_heading(doc, "5.3 Llamadas de voz", 2)
    doc.add_paragraph(
        "Canal VOICE: orientado a llamadas telefónicas (integración prevista con agentes de voz "
        "y carrier). En el modelo de datos, comparte el mismo registro de eventos que los otros canales "
        "para unificar reporting."
    )
    add_heading(doc, "5.4 Estrategia de contacto: digital primero, voz como último recurso", 2)
    doc.add_paragraph(
        "La filosofia de diseño es: comenzar con canales de bajo costo y alta trazabilidad "
        "(WhatsApp y email) en ventanas definidas; reservar la llamada de voz para casos donde "
        "no hay respuesta o según políticas de propensión / severidad. "
        "Las secuencias permiten expresar esa prioridad (orden de pasos y límites por canal)."
    )

    # 6 IA
    add_heading(doc, "6. Agentes de inteligencia artificial", 1)
    doc.add_paragraph(
        "En mensajes entrantes de WhatsApp (y flujos relacionados), el sistema puede: "
        "obtener contexto del cliente, cargar políticas de cobranza, construir un prompt "
        "y generar una respuesta con modelo de lenguaje; luego enviar la respuesta por el mismo canal "
        "y registrar el evento saliente."
    )
    doc.add_paragraph(
        "Esto no reemplaza al operador en todos los casos: ante fallos de IA se puede enviar "
        "un mensaje de disculpa genérico y dejar el caso trazado para seguimiento humano."
    )

    # 7 Clasificación comportamiento
    add_heading(doc, "7. Comportamiento e historial de pago del cliente", 1)
    doc.add_paragraph(
        "La información operativa se apoya en: estado de facturas (abierta, vencida, parcial, saldada, etc.), "
        "aplicaciones de pago (payment_applications), promesas (promises) y eventos de contacto."
    )
    doc.add_paragraph(
        "A partir de mensajes y correos, el sistema puede extraer intenciones (pago, consulta, reclamo), "
        "montos o fechas mencionadas, y crear promesas o callbacks programados automáticamente "
        "cuando el texto lo permite (procesamiento de callbacks desde mensaje)."
    )

    # 8 Callbacks
    add_heading(doc, "8. Generación automática de callbacks y seguimiento", 1)
    doc.add_paragraph(
        "Cuando un cliente deja compromisos o fechas en texto o voz, el motor puede crear "
        "scheduled_callbacks (recordatorios de seguimiento) vinculados al cliente y, si aplica, a la factura."
    )
    doc.add_paragraph(
        "Esto reduce trabajo manual del operador y mantiene una cadena de seguimiento medible."
    )

    # 9 Cobranzas
    add_heading(doc, "9. Gestión de cobranzas", 1)
    doc.add_paragraph(
        "Cobranza en Constanza cubre: cartera de clientes y facturas, comunicación multicanal, "
        "promesas, KPIs, lotes de notificación y decisiones operativas cuando hay ambigüedad o conflicto."
    )
    add_heading(doc, "9.1 Políticas de cobranza", 2)
    doc.add_paragraph(
        "Las políticas (negociaciones máximas, plazos, porcentaje mínimo de pago, horarios de contacto, tono) "
        "pueden almacenarse por tenant y consumirse tanto por la UI/agente como por los prompts de IA."
    )

    # 10 Conciliación
    add_heading(doc, "10. Conciliación e ingreso de pagos (backend)", 1)
    doc.add_paragraph(
        "Los pagos se modelan en pay.payments con método de ingreso (por ejemplo transferencia), "
        "origen (CRESIUM, MANUAL, etc.), estado y vínculo a application por factura."
    )
    add_heading(doc, "10.1 Webhooks Cresium (rail-cucuru)", 2)
    doc.add_paragraph(
        "El servicio rail-cucuru recibe webhooks de depósitos/transferencias, valida firma HMAC cuando "
        "corresponde, aplica idempotencia y persiste el pago. Puede intentar imputación automática a factura "
        "usando número de factura en el payload, CUIT/CVU del cliente en cartera, o reglas opcionales "
        "(por ejemplo match solo por monto si está habilitado explícitamente)."
    )
    doc.add_paragraph(
        "Tras imputación, se sincroniza el estado de la factura según aplicaciones de pago."
    )
    add_heading(doc, "10.2 Notificaciones post-pago", 2)
    doc.add_paragraph(
        "Puede configurarse notificación por WhatsApp al receptor interno y, si hay imputación y teléfono "
        "del cliente, confirmación al pagador. El envío se canaliza a través del notifier y "
        "variables de entorno (URL del notifier, teléfono receptor, nombre comercial)."
    )

    # 11 Opciones
    add_heading(doc, "11. Opciones y configuración relevantes", 1)
    doc.add_paragraph(
        "A nivel tenant: CVU de cobro Cresium, datos de clientes y facturas, políticas de cobranza."
    )
    doc.add_paragraph(
        "A nivel integración: secretos de webhook Cresium, tenant destino de depósitos, flags de rechazo "
        "por CVU, unidad de monto (pesos vs centavos), ventana anti-replay de timestamps."
    )
    doc.add_paragraph(
        "A nivel notifier: Redis, BuilderBot (API key, bot id, base URL), SMTP para email."
    )
    doc.add_paragraph(
        "A nivel API: JWT, CORS, URL del notifier para envíos desde la aplicación."
    )

    # 12 Cierre
    add_heading(doc, "12. Resumen ejecutivo", 1)
    doc.add_paragraph(
        "Constanza integra contact center omnicanal, contexto unificado por cliente y factura, "
        "IA asistida para respuestas de cobranza, extracción de compromisos y callbacks, "
        "y motor de pagos con conciliación automática vía webhooks — con trazabilidad end-to-end "
        "desde el primer contacto hasta el pago imputado."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Nota: Algunos microservicios de la arquitectura global (orquestador de contacto dedicado, "
        "reconciler standalone, cola de decisiones como servicio separado) pueden estar en distinto grado "
        "de madurez; la base de datos y el notifier ya soportan el modelo de eventos y canales descrito."
    ).italic = True

    out = "/Users/ralborta/Constanza/docs/Constanza-Funcionalidad-Contact-Center-y-Cobranzas.docx"
    doc.save(out)
    print(f"OK: {out}")

if __name__ == "__main__":
    main()
